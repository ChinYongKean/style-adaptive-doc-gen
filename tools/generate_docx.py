#!/usr/bin/env python3
"""
Style-Adaptive DOCX Generator
Reads a YAML style profile + markdown content → produces styled .docx

Usage:
    python generate_docx.py --profile examples/style-profiles/enterprise-cloud-consulting.yaml \
                            --content input.md \
                            --output report.docx
"""
import argparse
import re
import yaml
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color.lstrip("#")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="BFBFBF"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def load_profile(path):
    with open(path) as f:
        return yaml.safe_load(f)


def parse_markdown(md_text):
    blocks = []
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        # Heading
        m = re.match(r'^(#{1,5})\s+(.+)$', line)
        if m:
            blocks.append(('heading', len(m.group(1)), m.group(2).strip()))
            i += 1
            continue
        # Table
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers = [c.strip() for c in line.split('|')[1:-1]]
            i += 2
            rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                rows.append([c.strip() for c in lines[i].split('|')[1:-1]])
                i += 1
            blocks.append(('table', headers, rows))
            continue
        # Bullet
        m = re.match(r'^[-*+]\s+(.+)$', line)
        if m:
            blocks.append(('bullet', m.group(1).strip()))
            i += 1
            continue
        # Paragraph
        if line.strip():
            blocks.append(('paragraph', line.strip()))
        i += 1
    return blocks


def generate(profile, content_md, output_path):
    p = profile
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = p['brand']['body_font']
    style.font.size = Pt(p['body']['font_size_pt'])
    style.font.color.rgb = hex_to_rgb(p['body']['color'])

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(p['document']['page_margins']['top_cm'])
        section.bottom_margin = Cm(p['document']['page_margins']['bottom_cm'])
        section.left_margin = Cm(p['document']['page_margins']['left_cm'])
        section.right_margin = Cm(p['document']['page_margins']['right_cm'])

    # Parse and render content
    blocks = parse_markdown(content_md)

    for block in blocks:
        btype = block[0]

        if btype == 'heading':
            _, level, text = block
            heading = doc.add_heading(text, level=min(level, 3))
            run = heading.runs[0]
            run.font.name = p['brand']['heading_font']
            hkey = f'h{min(level, 3)}'
            run.font.size = Pt(p['headings'][hkey]['size_pt'])
            run.font.color.rgb = hex_to_rgb(p['headings'][hkey]['color'])
            run.font.bold = p['headings'][hkey]['bold']

        elif btype == 'table':
            _, headers, rows = block
            table = doc.add_table(rows=0, cols=len(headers))
            # Header row
            hdr_row = table.add_row()
            for i, h in enumerate(headers):
                cell = hdr_row.cells[i]
                cell.text = ''
                run = cell.paragraphs[0].add_run(h)
                run.font.name = p['brand']['heading_font']
                run.font.size = Pt(10)
                run.font.bold = p['tables']['header_bold']
                run.font.color.rgb = hex_to_rgb(p['tables']['header_text_color'])
                set_cell_shading(cell, p['tables']['header_bg'])
                set_cell_borders(cell, p['tables']['header_bg'].lstrip('#'))
            # Data rows
            for row_data in rows:
                row = table.add_row()
                for i, val in enumerate(row_data):
                    if i < len(row.cells):
                        cell = row.cells[i]
                        cell.text = ''
                        run = cell.paragraphs[0].add_run(val)
                        run.font.name = p['brand']['body_font']
                        run.font.size = Pt(10)
                        set_cell_borders(cell)
            doc.add_paragraph()

        elif btype == 'bullet':
            text = block[1]
            para = doc.add_paragraph(style='List Bullet')
            para.clear()
            # Handle inline bold
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = para.add_run(part)
                run.font.name = p['brand']['body_font']
                run.font.size = Pt(p['body']['font_size_pt'])

        elif btype == 'paragraph':
            text = block[1]
            para = doc.add_paragraph()
            align_map = {'justified': WD_ALIGN_PARAGRAPH.JUSTIFY, 'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER}
            para.alignment = align_map.get(p['body']['alignment'], WD_ALIGN_PARAGRAPH.LEFT)
            # Handle inline bold/italic
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = para.add_run(part[1:-1])
                    run.font.italic = True
                else:
                    run = para.add_run(part)
                run.font.name = p['brand']['body_font']
                run.font.size = Pt(p['body']['font_size_pt'])
                run.font.color.rgb = hex_to_rgb(p['body']['color'])

    doc.save(output_path)
    print(f"[OK] Generated: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Generate styled DOCX from profile + markdown")
    parser.add_argument("--profile", required=True, help="YAML style profile path")
    parser.add_argument("--content", required=True, help="Markdown content file")
    parser.add_argument("--output", "-o", required=True, help="Output .docx path")
    args = parser.parse_args()

    profile = load_profile(args.profile)
    content = Path(args.content).read_text('utf-8')
    generate(profile, content, args.output)


if __name__ == "__main__":
    main()
